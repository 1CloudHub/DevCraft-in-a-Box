import pandas as pd
import re
from urllib.parse import urlparse
from botocore.exceptions import ClientError  
from bs4 import BeautifulSoup
import boto3, json
import warnings
from io import StringIO, BytesIO  
import base64, random ,string
from datetime import *
import time                                   
from decimal import Decimal    
import os
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table  
import sys  
import psycopg2


schema = os.environ["schema"]
region_used = os.environ["region_used"]
bucket_name = os.environ["bucket_name"]  
db_database = os.environ["db_database"]
db_password = os.environ["db_password"]
db_host = os.environ["db_host"]
db_port = os.environ["db_port"]
db_user = os.environ["db_user"]
file_version_table = os.environ["file_version_table"]
file_metadata_Table = os.environ["file_metadata_table"]
app_metadata_table = os.environ["app_metadata_table"]
preprocess_table = os.environ["preprocess_table"]
model_id_sonnet = os.environ["model_id_sonnet"]
file_log_table = os.environ["file_log_table"]

Document_Processing = os.environ['Document_Processing']
Injestion_trigger = os.environ['Injestion_trigger']


s3_client = boto3.client('s3',region_name = region_used)
s3_resource = boto3.resource('s3',region_name = region_used) 



def db_select(query): #select
    try:
        connection = psycopg2.connect(
            user=db_user,
            password=db_password,
            host=db_host,
            port=db_port,  # Replace with the SSH tunnel local bind port
            database=db_database
        )                       
        cursor = connection.cursor()
        cursor.execute(query)
        result = cursor.fetchall()
        connection.commit()
        cursor.close()
        connection.close()
        return result
    except Exception as e:
        result = []
        re = {
            "statusCode": 500,
            "status" : "failed",
            "message": f"The following exception occurred: {str(e)}"
        }
        result.append(re)
        return result
        
def update_db(query): #update
    try:
        connection = psycopg2.connect(
            user=db_user,
            password=db_password,
            host=db_host,
            port=db_port,  
            database=db_database
        )                                  
        cursor = connection.cursor()
        cursor.execute(query)
        connection.commit()
        cursor.close()
        connection.close()  
        return {"status" : "success"}
    except Exception as e:
        result = []
        re = {
            "statusCode": 500,
            "message": f"The following exception occurred: {str(e)}"
        }
        result.append(re)
        print(result)
        return {"status" : "failed"}
        
def db_insert(query): #insert
    try:
        connection = psycopg2.connect(
            user=db_user,
            password=db_password,
            host=db_host,
            port=db_port,  # Replace with the SSH tunnel local bind port
            database=db_database
        )    
        cursor = connection.cursor()
        cursor.execute(query)
        connection.commit()
        cursor.close()
        connection.close()
        return {"status" : "success"}
    except Exception as e:
        result = []
        re = {
            "statusCode": 500,
            "message": f"The following exception occurred: {str(e)}"
        }
        result.append(re)
        print(result)
        return {"status" : "failed"}
        
    
#MOVE FILE FROM STAGING FOLDER TO MAIN FOLDER
def move_staging_file(bucket_name,input_file_key,destination_key,app_id,env):
    try:
        # Define destination and source paths (swapped)
        # app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        # app_name_result = db_select(app_query)
        # app_name = app_name_result[0][0] 

        # Copy the file from staging folder to kb_folder and 
        print("MOVING FROM STAGING FOLDER")
        print("DESTINATION KEY : ",destination_key)
        print("INPUT KEY : ",input_file_key)
        response = s3_resource.Object(bucket_name, destination_key).copy_from(CopySource={'Bucket': bucket_name, 'Key': input_file_key})
        print("MOVE RESPONSE : ",response)
        # Delete the file from the current location
        response = s3_client.delete_object(Bucket=bucket_name, Key=input_file_key)
        print("DELETE RESPONSE : ",response)
        print(f"Moved {input_file_key} to {destination_key}")
        return f"Moved {input_file_key} to {destination_key} successfully"
    except Exception as e:
        print("An exception occurred at function move_staging_file:",e)
        return "Moving staging file failed"

# MOVE S3 FOLDER FUNCTION
def move_s3_folder(bucket_name, actual_file_name, new_version,file_to_move,app_id,env):  
    try:
        # Define source and destination paths
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"    
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        source_prefix = f"{app_name}/Dev_Documents_kb/{file_to_move}/"
        print("MOVE FROM SOURCE : ",source_prefix)
        destination_prefix = f"{app_name}/Dev_Documents/{actual_file_name}_versions/{actual_file_name}_version{new_version}/preprocessed_files/{file_to_move}/"
        print("MOVE TO DEST : ",destination_prefix)
        # List all objects in the source folder
        paginator = s3_client.get_paginator('list_objects_v2')
        page_iterator = paginator.paginate(Bucket=bucket_name,  Prefix=source_prefix)
        for page in page_iterator: 
            if 'Contents' in page:
                for obj in page['Contents']:
                    # Get the source key
                    source_key = obj['Key']
                    
                    # Define the destination key
                    destination_key = destination_prefix + source_key[len(source_prefix):]
                    
                    # Copy the object to the new location
                    s3_resource.Object(bucket_name, destination_key).copy_from(CopySource={'Bucket': bucket_name, 'Key': source_key})
                    
                    # Delete the original object
                    s3_client.delete_object(Bucket=bucket_name, Key=source_key)
                    print(f"Moved {source_key} to {destination_key}")
    except Exception as e:
        print("An exception occurred at function move_s3_folder: ",e)
        
# MOVE S3 WHOLE FILE FUNCTION
def move_s3_orginal_file(bucket_name,actual_file_name, new_version, file_to_move,app_id,env):
    try:
        # Define source and destination paths
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        source_key = f"{app_name}/Dev_Documents_kbview/{file_to_move}"
        print("MOVE FROM SOURCE : ",source_key)
        destination_key = f"{app_name}/Dev_Documents/{actual_file_name}_versions/{actual_file_name}_version{new_version}/whole_files/{file_to_move}"
        print("MOVE FROM DEST : ",destination_key)
        # Copy the file to the new location
        s3_resource.Object(bucket_name, destination_key).copy_from(CopySource={'Bucket': bucket_name, 'Key': source_key})
        
        # Delete the original file
        s3_client.delete_object(Bucket=bucket_name, Key=source_key)
        print(f"Moved {source_key} to {destination_key}")
    except Exception as e:
        print("An exception occurred at function move_s3_orginal_file: ",e)

def move_s3_file(bucket_name, actual_filename, new_version, file_to_move,app_id,env):
    print("IN S3 FUNNCTION TO COPY KB FOLDER")
    # Construct the source and destination keys for the main file
    app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
    app_name_result = db_select(app_query)
    app_name = app_name_result[0][0]
    source_key = f"{app_name}/Dev_Documents_kb/{file_to_move}"
    print("MOVE FROM SOURCE : ",source_key)
    destination_key = f"{app_name}/Dev_Documents/{actual_filename}_versions/{actual_filename}_version{new_version}/{file_to_move}"
    print("MOVE TO DEST : ",destination_key)
    # Construct the source and destination keys for the metadata file
    metadata_source_key = f"{app_name}/Dev_Documents_kb/{file_to_move}.metadata.json"
    metadata_destination_key = f"{app_name}/Dev_Documents/{actual_filename}_versions/{actual_filename}_version{new_version}/{file_to_move}.metadata.json"

    try:
        # Copy the main file to the new location
        copy_source = {
            'Bucket': bucket_name,
            'Key': source_key
        }
        s3_client.copy(copy_source, bucket_name, destination_key)
        print(f"Successfully copied {source_key} to {destination_key}")

        # Delete the original main file
        s3_client.delete_object(Bucket=bucket_name, Key=source_key)
        print(f"Successfully deleted {source_key}")

        # Copy the metadata file to the new location
        metadata_copy_source = {
            'Bucket': bucket_name,
            'Key': metadata_source_key
        }
        s3_client.copy(metadata_copy_source, bucket_name, metadata_destination_key)
        print(f"Successfully copied {metadata_source_key} to {metadata_destination_key}")

        # Delete the original metadata file
        s3_client.delete_object(Bucket=bucket_name, Key=metadata_source_key)
        print(f"Successfully deleted {metadata_source_key}")
    except ClientError as e:
        print(f"Error occurred: {e}")
        raise e 
        
# MOVE S3 FOLDER REVERSE FUNCTIONS

def move_s3_folder_reverse(bucket_name,actual_file_name, new_version, file_to_move,app_id,env):
    try:
        # Define destination and source paths (swapped)
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        destination_prefix = f"{app_name}/Dev_Documents_kb/{file_to_move}/"    
        print("MOVE TO DEST : ",destination_prefix)
        source_prefix = f"{app_name}/Dev_Documents/{actual_file_name}_versions/{actual_file_name}_version{new_version}/preprocessed_files/{file_to_move}/"
        print("MOVE FROM SOURCE : ",source_prefix)
        # List all objects in the destination folder (which was the source previously)
        paginator = s3_client.get_paginator('list_objects_v2')
        page_iterator = paginator.paginate(Bucket=bucket_name,  Prefix=source_prefix)
        for page in page_iterator:
            if 'Contents' in page:
                for obj in page['Contents']:
                    # Get the source key
                    source_key = obj['Key']
                    
                    # Define the destination key
                    destination_key = destination_prefix + source_key[len(source_prefix):]
                    
                    # Copy the object back to the original location
                    s3_resource.Object(bucket_name, destination_key).copy_from(CopySource={'Bucket': bucket_name, 'Key': source_key})
                    
                    # Delete the object from the current location
                    s3_client.delete_object(Bucket=bucket_name, Key=source_key)
                    print(f"Moved {source_key} to {destination_key}")
    except Exception as e:
        print("An error occurred in the function move_s3_folder_reverse: ",e)

def move_s3_orginal_file_reverse(bucket_name,actual_file_name, new_version, file_to_move,app_id,env):
    try:
        # Define destination and source paths (swapped)
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        destination_key = f"{app_name}/Dev_Documents_kbview/{file_to_move}"
        print("MOVE TO DEST : ",destination_key)
        source_key = f"{app_name}/Dev_Documents/{actual_file_name}_versions/{actual_file_name}_version{new_version}/whole_files/{file_to_move}"
        print("MOVE FROM SOURCE : ",source_key)
        # Copy the file back to the original location
        s3_resource.Object(bucket_name, destination_key).copy_from(CopySource={'Bucket': bucket_name, 'Key': source_key})
        
        # Delete the file from the current location
        s3_client.delete_object(Bucket=bucket_name, Key=source_key)
        print(f"Moved {source_key} to {destination_key}")
    except Exception as e:
        print("An exception occurred at function move_s3_orginal_file_reverse:",e)
        

def move_s3_file_reverse(bucket_name, actual_filename, new_version, file_to_move,app_id,env):
    print("IN S3 FILE REVERSE")
    # Construct the reversed source and destination keys for the main file
    app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
    app_name_result = db_select(app_query)
    app_name = app_name_result[0][0]
    destination_key = f"{app_name}/Dev_Documents_kb/{file_to_move}"
    print("MOVE TO DEST : ",destination_key)
    source_key = f"{app_name}/Dev_Documents/{actual_filename}_versions/{actual_filename}_version{new_version}/{file_to_move}"
    print("MOVE FROM SOURCE : ",source_key)
    # Construct the reversed source and destination keys for the metadata file
    metadata_destination_key = f"{app_name}/Dev_Documents_kb/{file_to_move}.metadata.json"
    metadata_source_key = f"{app_name}/Dev_Documents/{actual_filename}_versions/{actual_filename}_version{new_version}/{file_to_move}.metadata.json"

    try:
        # Copy the main file to the new location (reversed keys)
        copy_source = {
            'Bucket': bucket_name,
            'Key': source_key
        }
        s3_client.copy(copy_source, bucket_name, destination_key)
        print(f"Successfully copied {source_key} to {destination_key}")

        # Delete the original main file (reversed keys)
        s3_client.delete_object(Bucket=bucket_name, Key=source_key)
        print(f"Successfully deleted {source_key}")

        # Copy the metadata file to the new location (reversed keys)
        metadata_copy_source = {
            'Bucket': bucket_name,  
            'Key': metadata_source_key
        }
        s3_client.copy(metadata_copy_source, bucket_name, metadata_destination_key)
        print(f"Successfully copied {metadata_source_key} to {metadata_destination_key}")

        # Delete the original metadata file (reversed keys)
        s3_client.delete_object(Bucket=bucket_name, Key=metadata_source_key)
        print(f"Successfully deleted {metadata_source_key}")        

    except ClientError as e:
        print(f"Error occurred: {e}")
        raise e


# NEW CSV UPLOAD FUNCTION
def new_csv_upload_process(file_name,file_type,access_type,uploaded_by,input_file_key,app_id,env):
    try:
        file_id = None
        #checking for file name:
        query = f"""
                SELECT * FROM {schema}.{file_version_table}
                WHERE active_file_name = '{file_name}' AND delete_status = 0 AND app_id = '{app_id}' and env = '{env}';
                """  
        print("CHECKING IF FILE EXISTS")
        response = db_select(query)
        print("FILE NAME CHECK RDS RESPONSE",response)
        if len(response) != 0:
            return {"status":"500","Message":"FILE NAME ALREADY EXISTS"}
        print("FILE DOES NOT EXIST NEW FILE INSERTING NOW")
        #insert into  file_metadata
        file_id= ''.join(random.choices(string.ascii_uppercase + string.digits, k=10))
        current_time = datetime.now(timezone.utc).isoformat()
        query = f"""
            INSERT INTO {schema}.{file_metadata_Table} (
                id,uploaded_by,actual_filename, total_version, active_version, active_file_name, 
                sync_status, access_type, last_updated_time, delete_status, file_type, remain_count_kb,trigger_count,app_id,env
            ) VALUES (
                '{file_id}','{uploaded_by}' , '{file_name}', 1, 1, '{file_name}', 
                'INPROGRESS', '{access_type}', '{current_time}', 0, '{file_type}', 0, 0,'{app_id}','{env}'
            );
            """
        put_response = db_insert(query)
        print("DB RESPONSE AFTER INSRTING : ",put_response)
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0] 
        object_key  = f"{app_name}/Dev_Documents_kbview/{file_name}"
        print("OBJECT KEY",object_key)
        #moving file from staging to kb_view

        # s3_upload = s3_client.put_object(Bucket=bucket_name, Key=object_key, Body=content_decoded) 
        s3_upload = move_staging_file(bucket_name,input_file_key,object_key,app_id,env)
        print("S3 RESPONSE : ",s3_upload)
        lambda_client = boto3.client('lambda',region_name = region_used)
        lambda_response = lambda_client.invoke(
            FunctionName = Document_Processing,                                                                                
            Payload = json.dumps({
                'event_type' :'csv_file_process',
                'file_name' : file_name,
                'file_type' : file_type,
                'object_key' : object_key,
                "access_type" : access_type,
                'uploaded_by' : uploaded_by,
                "file_id" : file_id,
                "app_id" : app_id,
                "app_name"  :app_name,
                "env" : env
                
            }),   
            InvocationType = 'Event'
            )
        query_insert = f"""
        INSERT INTO {schema}.{file_log_table} (id, actual_file_name, active_file_name, done_by, action,app_id,env)
        VALUES (
            '{file_id}',
            '{file_name}',
            '{file_name}',
            '{uploaded_by}',
            'New File Addedd',
            '{app_id}',
            '{env}'
        );
        """
        
        # Execute the INSERT query using the existing db_insert function
        db_insert(query_insert)
        return {"status":"200","Message":"Upload Successful"}
    except Exception as e:
        print("First exception")   
        print("Error processing file : ",e)
        if file_id:
            query = f"""
            UPDATE {schema}.{file_metadata_Table}
            SET sync_status = 'FAILED'
            WHERE id = '{file_id}';
            """

            # Execute the query
            update_response = execute_query(query)
        return {"status":"500","Message":"Upload failed","error":str(e)}   
        
        
# NEW FILE UPLOAD FUNCTION {NOT CSV} DIRECT UPLOAD
def new_file_upload_process(file_name,file_type,access_type,uploaded_by,input_file_key,app_id,env):
    try:
        file_id = None
        #checking for file name:
        query = f"""
                SELECT * FROM {schema}.{file_version_table}
                WHERE active_file_name = '{file_name}' AND delete_status = 0 AND app_id = '{app_id}' AND env = '{env}';
                """  
        print("CHECKING IF FILE EXISTS")
        response = db_select(query)
        print("FILE NAME CHECK RDS RESPONSE",response)
        if len(response) != 0:
            return {"status":"500","Message":"FILE NAME ALREADY EXISTS"}
        print("FILE DOES NOT EXIST NEW FILE INSERTING NOW")      
        
        #insert into dynamo_db file_metadata
        current_time = datetime.now(timezone.utc).isoformat()
        file_id= ''.join(random.choices(string.ascii_uppercase + string.digits, k=10))
        query = f"""
        INSERT INTO {schema}.{file_metadata_Table} (
            id,uploaded_by,actual_filename, total_version, active_version, active_file_name,
            sync_status, access_type, last_updated_time, delete_status, file_type, count_in_kb, remain_count_kb,trigger_count,app_id,env
        ) VALUES (
            '{file_id}','{uploaded_by}', '{file_name}', 1, 1, '{file_name}',
            'INPROGRESS', '{access_type}', '{current_time}', 0, '{file_type}', 1, 0, 0,'{app_id}','{env}'
        );
        """
        response = db_insert(query)
        
        #insert into dynamo_db file_versions
        file_version_id = ''.join(random.choices(string.ascii_uppercase + string.digits, k=11))
        query = f"""
        INSERT INTO {schema}.{file_version_table} (
            id, actual_file_name, active_file_name, version,
            access_type, last_updated_time, delete_status, uploaded_by, count_in_kb,app_id,env
        ) VALUES (
            '{file_version_id}','{file_name}', '{file_name}', 1,
            '{access_type}', '{current_time}', 0, '{uploaded_by}', '1','{app_id}','{env}'
        );
        """
        
        # Execute the insert query
        response = db_insert(query)
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        object_key  = f"{app_name}/Dev_Documents_kb/{file_name}"
        # s3_upload = s3_client.put_object(Bucket=bucket_name, Key=object_key, Body=content_decoded)
        s3_upload = move_staging_file(bucket_name,input_file_key,object_key,app_id,env)
        print("S3 RESPONSE : ",s3_upload)
        if access_type == 'private':
            metadata = {
            "Access":"private",
            }
            
        else:
            metadata = {
            "Access":"public",
            }
        if file_type == 'csv':
            try:
                csv_content = s3_client.get_object(Bucket=bucket_name, Key=object_key)['Body'].read().decode('latin1')
                print(csv_content)
                df = pd.read_csv(StringIO(csv_content))
        
            except Exception as e:
                print(e)
                if file_id:
                    table_name = file_metadata_Table
                    update_db(table_name,'id',file_id,'sync_status','FAILED')
                return {"status":"500","Message":"Upload failed","error":e}
                        
        filtered_metadata = {k: v for k, v in metadata.items() if v and not v.isspace()}
        metadata_dict = {"metadataAttributes": filtered_metadata}
        metadata_content = json.dumps(metadata_dict, indent=4)
        object_key_metadata  = f"{app_name}/Dev_Documents_kb/{file_name}.metadata.json"
        s3_upload = s3_client.put_object(Bucket=bucket_name, Key=object_key_metadata, Body=metadata_content, ContentType='application/json')
        print(f"Metadata file uploaded successfully ") 
        query_insert = f"""
        INSERT INTO {schema}.{file_log_table} (id, actual_file_name, active_file_name, done_by, action,app_id,env)
        VALUES (
            '{file_id}',
            '{file_name}',
            '{file_name}',
            '{uploaded_by}',
            'New File Addedd',
            '{app_id}',
            '{env}'
        );
        """
        # Execute the INSERT query using the existing db_insert function
        db_insert(query_insert)
        
        lambda_client = boto3.client('lambda',region_name = region_used)
        lambda_response = lambda_client.invoke(
            FunctionName = Injestion_trigger,                                                                                
            Payload = json.dumps({
                'file_name':file_name,
                'bucket_name':bucket_name,
                'trigger_action':'Put',
                'app_name':app_name
            }),   
            InvocationType = 'Event'
        )
        
        return {"status":"200","Message":"Upload successful"}    
    except Exception as e:
        print("Exception occurred:",e)
        if file_id:
            query = f"""
            UPDATE {schema}.{file_metadata_Table}
            SET sync_status = 'FAILED'
            WHERE id = '{file_id}';
            """
            response = update_db(query)
        return {"status":"500","Message":"Upload failed","error":e}
        
# NEW CSV VERSION UPLOAD VERSION 
def new_version_csv_upload_process(actual_file_name,file_name,file_type,access_type,uploaded_by,input_file_key,app_id,env):
    try:
        id_to_update = None
        #checking for file name:
        query = f"""
                SELECT * FROM {schema}.{file_version_table}
                WHERE active_file_name = '{file_name}' AND delete_status = 0 AND app_id = '{app_id}' AND env = '{env}';
                """  
        print("CHECKING IF FILE EXISTS")
        response = db_select(query)
        print("FILE NAME CHECK RDS RESPONSE",response)
        if len(response) != 0:
            return {"status":"500","Message":"FILE NAME ALREADY EXISTS"}
        print("FILE DOES NOT EXIST NEW FILE INSERTING NOW")     
        
        #Checking for inprogress:
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE active_file_name = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'INPROGRESS'
          AND app_id = '{app_id}' AND env = '{env}';
        """
        
        result = db_select(query)
        print("IN PROGRESS QUERY RESULT",result)
        
        if result:
            return {"status": "100", "Message": "FILE IS INPROGRESS"}
        
        #Checking for deleting
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'DELETING'
          AND app_id = '{app_id}' AND env = '{env}';
        """
        
        result = db_select(query)
        print("DELETING QUERY RESULT : ",result)

        if result:
            return {"status": "100", "Message": "FILE IS DELETING"}
        
        #Scan data by giving actual_file_name 
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND app_id = '{app_id}' AND env = '{env}';
        """
        result = db_select(query)
        print("SCANNING RESULT : ",result)
        if result:
            # Extracting values from the first item (assuming only one item is returned)
            version_active = int(result[0][4])
            total_version = int(result[0][12])
            new_version = total_version + 1
            id_to_update = result[0][0]
            previous_active_file_name = result[0][2]
            current_time = datetime.now(timezone.utc).isoformat()
        #update total_versions, active_version, active_file_name
        query = f"""
        UPDATE {schema}.{file_metadata_Table}
        SET 
            sync_status = 'INPROGRESS',
            total_version = '{new_version}',
            active_version = '{new_version}',
            active_file_name = '{file_name}',
            last_updated_time = '{current_time}',
            remain_count_kb = 0,
            trigger_count = 0
        WHERE id = '{id_to_update}';
        """
        
        response = update_db(query)
        print(" RESPONSE ACTIVE FILE NAME UPDATED IN FILE METADATA TABLE : ",response)
        
        
        move_s3_folder(bucket_name, actual_file_name, int(version_active), previous_active_file_name,app_id,env)
        time.sleep(3)
        move_s3_orginal_file(bucket_name,actual_file_name,int(version_active), previous_active_file_name,app_id,env)
        
        #insert new file
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        object_key  = f"{app_name}/Dev_Documents_kbview/{file_name}"
        #moving staging file
        s3_upload = move_staging_file(bucket_name,input_file_key,object_key,app_id,env)
        print("S3 RESPONSE : ",s3_upload)
        # s3_upload = s3_client.put_object(Bucket=bucket_name, Key=object_key, Body=content_decoded)
        lambda_client = boto3.client('lambda',region_name = region_used)
        lambda_response = lambda_client.invoke(
            FunctionName = Document_Processing,                                                                                
            Payload = json.dumps({
                'event_type' :'new_csv_version_processing',
                "actual_file_name" : actual_file_name,
                "new_version" : new_version,
                'file_name' : file_name,
                'file_type' : file_type,
                'object_key' : object_key,
                'access_type' : access_type,
                "id_to_update" : id_to_update,
                'uploaded_by' : uploaded_by,
                "app_id" : app_id,
                "app_name" : app_name,
                "env" : env
                
                
            }),   
            InvocationType = 'Event'
            )
        
        return {"status":"200","Message":"Upload Successful"}
    except Exception as e:
        print("An Exception occurred:",e)
        if id_to_update:
            query = f"""
            UPDATE {schema}.{file_metadata_Table}
            SET sync_status = 'FAILED'
            WHERE id = '{id_to_update}';
            """
            response = update_db(query)
        return {"status":"500","Message":"Upload failed","error":e}
        
# NEW VERSION NORMAL FILE UPLOAD 
def new_version_file_upload_process(actual_file_name,file_name,file_type,access_type,uploaded_by,input_file_key,app_id,env):
    try:
        id_to_update = None
        #checking for file name:
        query = f"""
        SELECT * FROM {schema}.{file_version_table}
        WHERE active_file_name = '{file_name}' AND delete_status = 0 AND app_id = '{app_id}' AND env = '{env}';
        """  
        print("CHECKING IF FILE EXISTS")
        response = db_select(query)
        print("FILE NAME CHECK RDS RESPONSE",response)
        if len(response) != 0:
            return {"status":"500","Message":"FILE NAME ALREADY EXISTS"}
        print("FILE DOES NOT EXIST NEW FILE INSERTING NOW")     
        
        #Checking for inprogress:
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE active_file_name = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'INPROGRESS'
          AND app_id = '{app_id}' AND env = '{env}';
        """
        result = db_select(query)
        
        if result:
            return {"status": "100", "Message": "FILE IS INPROGRESS"}
        
        #Checking for deleting
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'DELETING'
          AND app_id = '{app_id}' AND env = '{env}';
        """
        result = db_select(query)
        if result:
            return {"status": "100", "Message": "FILE IS DELETING"}
    
        #Scan data by giving actual_file_name 
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND app_id = '{app_id}' AND env = '{env}';
        """
        result = db_select(query)
        if result:
            # Extracting values from the first item (assuming only one item is returned)
            version_active = int(result[0][4])
            total_version = int(result[0][12])
            new_version = total_version + 1
            id_to_update = result[0][0]
            previous_active_file_name = result[0][2]
            current_time = datetime.now(timezone.utc).isoformat()       
        #update total_versions, active_version, active_file_name
        query = f"""
        UPDATE {schema}.{file_metadata_Table}
        SET 
            sync_status = 'INPROGRESS',
            total_version = '{new_version}',
            active_version = '{new_version}',
            active_file_name = '{file_name}',
            last_updated_time = '{current_time}',
            remain_count_kb = 0,
            trigger_count = 0
        WHERE id = '{id_to_update}';
        """
        
        response = update_db(query)
       
        #insert new version into file_version table
        file_version_id = ''.join(random.choices(string.ascii_uppercase + string.digits, k=11))
        query_insert = f"""
        INSERT INTO {schema}.{file_version_table} (
            id, actual_file_name, active_file_name, version, access_type, last_updated_time, delete_status, uploaded_by, count_in_kb,app_id,env
        ) VALUES (
            '{file_version_id}','{actual_file_name}', '{file_name}', {new_version}, '{access_type}', '{current_time}', 0, '{uploaded_by}', 1,'{app_id}','{env}'
        );
        """
        
        # Execute the INSERT query
        put_response = db_insert(query_insert)
        
        #Move the file from kb folder 
        move_s3_file(bucket_name, actual_file_name, int(version_active), previous_active_file_name,app_id,env)   
        #Insert the new file
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        object_key  = f"{app_name}/Dev_Documents_kb/{file_name}"  
        # s3_upload = s3_client.put_object(Bucket=bucket_name, Key=object_key, Body=content_decoded)
        s3_upload = move_staging_file(bucket_name,input_file_key,object_key,app_id,env)
        print("S3 RESPONSE : ",s3_upload)
        metadata = {}
        # Upload Metadata file
        if access_type == 'private':
            metadata = {
            "Access":"private"
            }
            
        else:
            metadata = {
            "Access":"public"
            }
        if file_type == 'csv':
            try:
                csv_content = s3_client.get_object(Bucket=bucket_name, Key=object_key)['Body'].read().decode('latin1')
                print(csv_content)
                df = pd.read_csv(StringIO(csv_content))
            except Exception as e:
                print(e)
                        
        filtered_metadata = {k: v for k, v in metadata.items() if v and not v.isspace()}
        metadata_dict = {"metadataAttributes": filtered_metadata}  
        metadata_content = json.dumps(metadata_dict, indent=4)
        object_key_metadata  = f"{app_name}/Dev_Documents_kb/{file_name}.metadata.json"
        s3_upload = s3_client.put_object(Bucket=bucket_name, Key=object_key_metadata, Body=metadata_content, ContentType='application/json')
        print(f"Metadata file uploaded successfully ")
        query_insert = f"""
        INSERT INTO {schema}.{file_log_table} (id, actual_file_name, active_file_name, done_by, action)
        VALUES (
            '{file_version_id}',
            '{actual_file_name}',
            '{file_name}',
            '{uploaded_by}',
            'New File Version Addedd',
            '{app_id}',
            '{env}'
        );
        """
        
        # Execute the INSERT query using the existing db_insert function
        db_insert(query_insert)
        
        lambda_client = boto3.client('lambda',region_name = region_used)
        lambda_response = lambda_client.invoke(
            FunctionName = Injestion_trigger,                                                                                
            Payload = json.dumps({
                'file_name':file_name,
                'bucket_name':bucket_name,
                'trigger_action':'Put',
                'app_name':app_name
            }),   
            InvocationType = 'Event'
        )
        
        return {"status":"200","Message":"Upload successful. Ingestion started ..."}
    except Exception as e:
        print("An Exception occurred:",e)
        if id_to_update:
            query = f"""
            UPDATE {schema}.{file_metadata_Table}
            SET sync_status = 'FAILED'
            WHERE id = '{id_to_update}';
            """
            response = update_db(query)
        return {"status":"500","Message":"Upload failed","error":e}


# WORD PRE PROCESSING FUNCTIONS START

# LLM WORD PRE PROCESSING FUNCTION
def llm_chain_for_nested_table(cur_data):
    boto3_bedrock = boto3.client(service_name='bedrock-runtime',region_name=region_used)
    print('cur_data',cur_data)
    TEMPLATE_FOR_NESTED_TABLE= f"""Human: We have information in the String format(contains several key-value pairs) for the <data></data> XML tags.
    <data>
    {cur_data}
    </data>
    Please convert this String information into a sentence.
    \nAssistant: """
    response = boto3_bedrock.invoke_model(contentType='application/json', body = json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 2000,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": TEMPLATE_FOR_NESTED_TABLE},
                    ]
                }
            ],
        })
        ,modelId=model_id_sonnet)
    inference_result = response['body'].read().decode('utf-8')
    final = json.loads(inference_result)
    final_response =  final['content'][0]['text']
    sys.stdout.flush()
    print(final_response)
    return final_response
 
# EXTRACTING ABOVE AND BELOW PARAGRAPH CONTENT
def get_nearby_text(table,num_paragraphs=10):
    nearby_paragraphs = []
   
    # Get paragraphs before the table
    previous_paragraphs = []
    for element in table._element.getprevious().iter():
        print("previous paragragh---------------------------------------------------")
        print(element)
        if element.tag.endswith('p'):
            previous_paragraphs.append(element)
    previous_paragraphs = previous_paragraphs[-num_paragraphs:]
 
    # Get paragraphs after the table
    next_paragraphs = []
    for element in table._element.getnext().iter():
        if element.tag.endswith('p'):
            print("next paragragh---------------------------------------------------")
            print(element)
            next_paragraphs.append(element)
    next_paragraphs = next_paragraphs[:num_paragraphs]
 
    # Extract text from nearby paragraphs
    for paragraph_element in previous_paragraphs + next_paragraphs:
        res = paragraph_element.text.strip().replace("✘", " no").replace("✔", " yes")
        nearby_paragraphs.append(res)
    print("Nearby paragraph-------------------------------------------------------------------------------")
    print(nearby_paragraphs)
    return nearby_paragraphs
 
# GET NESTED DATA FUNCTION 
def get_nested_data(row, table_title):
    row_text = f":"
    for cell in row.cells:
        # IF NO NESTED TABLE
        if not cell.tables:
            # do not have nested table
            row_text += cell.text.strip()
            row_text += ':'
            # continue
        # IF NESTED TABLE EXISTS   
        else:
            # exist nested table inside the cell
            cell_text = ''
            for nested_table in cell.tables:
                nearby_text = get_nearby_text(nested_table)
                print(f"nearby_text: {nearby_text}")
                table_dic = convert_nested_data(nested_table)
                table_text = dict_to_key_value_pairs(table_dic,nearby_text,table_title)
                print(f"table_text after LLM :{table_text}")
                cell_text += str(nearby_text) + ":" + str(table_text)
                cell_text += '\n'
           
            row_text += cell_text
            row_text += ":"
    return row_text
 
# CONVERTING NESTED DATA FUNCTION
def convert_nested_data(table):
    print("TABLE : ",table)
    keys_list = []
    dic = {}
    for index, row in enumerate(table.rows):
        print("TABLE ROW : ",row)
        print('index',index)
        print("KEYS LIST : ",keys_list)
        if index == 0:
            # get keys list
            for cell in row.cells:
                if cell.text.strip():
                    keys_list.append(cell.text.strip())
            dic = {key: [] for index, key in enumerate(keys_list)}
        else:
            row_text = ""
            # get values of each key.
            for index, cell in enumerate(row.cells):
                print("CONVERT NESTED DATA INDEEX != 0 : INDEX = ",index)
                print("CELL VALUE : ",cell.text.strip())
                if cell.text.strip():
                    dic[keys_list[index]].append(cell.text.strip())
    return dic
    
# LLM CALL FUCNTION 
def dict_to_key_value_pairs(dic,nearby_paragraph,table_title):
    cur_str = f'{table_title}'+":"+f'{nearby_paragraph} :'
    row_text = ''
    for key, value in dic.items():
        cur_str += str(key) + ":" + str(value)
        cur_str += '\n'
    cur_str = cur_str.strip().replace("✘", " no").replace("✔", " yes")
    print(f'cur_str( nested data before LLM): {cur_str}')
    answer = llm_chain_for_nested_table(cur_str)
    pattern = r'<data>([\s\S]*?)</data>'
    flag = 0
    match = re.findall(pattern, answer)
    if len(match) == 0:
        flag = 1
        pattern = r':\n\n(.*)'
    match = re.search(pattern, answer)
   
    matched_text = ""
    if match and flag == 1:
        matched_text = match.group(1).strip()
    elif match and flag ==0:
        matched_text = match[0].strip()
    else:
        print("No match found.")
    row_text += matched_text
    row_text += '\n'
    return row_text

# WORD PRE PROCESSING FUNCTION
def word_pre_processing(doc,file_type,uploaded_by,file_name,app_id,env):
    file_id = None
    df_list = []
    reference_info_threshold = 20
    doc_title = ''
    heading_one = ''
    heading_two = ''
    chunk_content = ''
    table_title = ''
    paragraph_content = ""
    content_decoded = base64.b64decode(doc)
    document = Document(BytesIO(content_decoded))                                  
    print('word-pre-processing-content_decoded',content_decoded)  

    #Input file uploaded in s3
    app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
    app_name_result = db_select(app_query)
    app_name = app_name_result[0][0]
    object_key = f"{app_name}/Document_Preprocessing/input_files/{file_name}"
    s3_upload = s3_client.put_object(Bucket=bucket_name, Key=object_key, Body=content_decoded)
    
    #insert into dynamo_db file_metadata
    file_id= ''.join(random.choices(string.ascii_uppercase + string.digits, k=10))
    current_time = datetime.now(timezone.utc).isoformat()
    query = f"""
    INSERT INTO {schema}.{preprocess_table} (id, input_filename, output_filename, preprocess_status, last_updated_time, delete_status, file_type, uploaded_by, description,app_id,env)
    VALUES (
        '{file_id}',
        '{file_name}',
        'output_{file_name}.csv',
        'INPROGRESS',
        '{current_time}',
        0,
        '{file_type}',
        '{uploaded_by}',
        'Successful',
        '{app_id}',
        '{env}'
    );
    """
    response = db_insert(query)
    try:
        for index, i in enumerate(document.iter_inner_content()):
            if isinstance(i,Paragraph): ## Process paragraph based on Headings
                if i.style.name == 'Title':
                    if doc_title != i.text.strip():
                        doc_title = i.text.strip()
                        print("DOC TITLE : ",doc_title)
                        continue   
                elif i.style.name == 'Heading 1':
                    if heading_one != i.text.strip():
                        print("PARAGRAPH CONTENT AFTER HEADING ! CHANGE AND BEFORE CLEARING : ",paragraph_content)
                        if paragraph_content != "":
                            df_list.append({'key':doc_title+":"+heading_one+":"+heading_two,'value':paragraph_content,'meta':doc_title})
                            paragraph_content = ""
                        heading_one = i.text.strip()
                        print("HEADING 1 : ",heading_one)
                        continue
                elif i.style.name == 'Heading 2':
                    if heading_two != i.text.strip():
                        if paragraph_content != "":
                            df_list.append({'key':doc_title+":"+heading_one+":"+heading_two,'value':paragraph_content,'meta':doc_title})
                            paragraph_content = ""              
                        heading_two = i.text.strip() 
                        continue
                if i.text.strip() != "":
                    text = i.text.strip() + ":\n"
                else:
                    continue
                paragraph_content += text
                continue

            if isinstance(i,Table):  ## Process Tables based on rows  
                table_title = heading_one + heading_two
                chunk_content = ''
                for index, row in enumerate(i.rows):                
                    inside_nested = any(cell.tables for cell in row.cells)
                    if inside_nested:
                        nested_data = get_nested_data(row,table_title) ## nested table processing
                        chunk_content += nested_data + '\n'
                    else:
                        for cell in row.cells:
                            chunk_content += '\n' + cell.text.strip() + '\n'
                if len(chunk_content)>0:
                    df_list.append({'key':doc_title+":"+heading_one+":"+heading_two,'value':chunk_content,'meta':doc_title})
                    continue
        if len(paragraph_content) > 0:
            print("APPENDING PARAGRAPH CONTENT IN DATAFRAME")
            df_list.append({'key':doc_title+":"+heading_one+":"+heading_two,'value':paragraph_content,'meta':doc_title})
            paragraph_content = ""
        df = pd.DataFrame(df_list)
        csv_file_content = df.to_csv(index=False)
        csv_file_name = f"output_{file_name}.csv"    
        csv_file_key = f"{app_name}/Document_Preprocessing/output_files/{csv_file_name}"
        put_output_file = s3_client.put_object(Bucket=bucket_name, Key=csv_file_key, Body=csv_file_content)
        query = f"""
        UPDATE {schema}.{preprocess_table}
        SET preprocess_status = 'COMPLETED'
        WHERE id = '{file_id}';
        """
        update_db(query)  
        query_insert = f"""
        INSERT INTO {schema}.{file_log_table} (id, actual_file_name, active_file_name, done_by, action,app_id,env)
        VALUES (
            '{file_id}',
            '{file_name}',
            '{file_name}',
            '{uploaded_by}',
            'Word pre processing',
            '{app_id}',
            'Production'
        );
        """
        
        # Execute the INSERT query using the existing db_insert function
        db_insert(query_insert)
        return {"status":"200","Message": "Upload successful"}
        
    except Exception as e:
        df = pd.DataFrame(df_list)
        csv_file_content = df.to_csv(index=False)
        csv_file_name = f"output_{file_name}.csv"    
        csv_file_key = f"Document_Preprocessing/output_files/{csv_file_name}"
        put_output_file = s3_client.put_object(Bucket=bucket_name, Key=csv_file_key, Body=csv_file_content)
        error_msg = "Errowr"+str(e) + "while  uploading occured at "+heading_one+":"+heading_two
        if file_id:
            query = f"""
            UPDATE {schema}.{preprocess_table}
            SET 
                preprocess_status = 'FAILED',
                description = '{error_msg}'
            WHERE id = '{file_id}';
            """
            
            response = update_db(query)
        return {"status":"500","Message": "Upload failed"}

# WORD PRE PROCESSING FUNCTIONS END


# FILE REVERTING FUNCTION START

def revert_csv_version(actual_file_name,new_file_name,new_version,done_by,app_id,env):
    try:
        id_to_update = None
        #Checking for inprogress:
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE active_file_name = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'INPROGRESS'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        result = db_select(query)
        
        if result:
            return {"status": "100", "Message": "FILE IS INPROGRESS"}
        
        #Checking for deleting
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'DELETING'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        result = db_select(query)
        if result:
            return {"status": "100", "Message": "FILE IS DELETING"}
        
        query_select = f"""
        SELECT * FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        
        # Execute the SELECT query using the existing db_select function
        response = db_select(query_select)
        existing_version = int(response[0][4])                                          
        existing_file_name = response[0][2]
        id_to_update = response[0][0]
        current_time = datetime.now(timezone.utc).isoformat()
        
        query_update = f"""
        UPDATE {schema}.{file_metadata_Table}
        SET sync_status = 'INPROGRESS',
            active_version = {new_version},
            active_file_name = '{new_file_name}',
            last_updated_time = '{current_time}',
            remain_count_kb = 0,
            trigger_count = 0
        WHERE id = '{id_to_update}';
        """
        
        # Execute the UPDATE query using the existing update_db function
        update_response = update_db(query_update)
        
        #file_versions scan

        print(actual_file_name)
        print(new_file_name)
        print(new_version)  
        move_s3_folder(bucket_name, actual_file_name, existing_version, existing_file_name,app_id,env)
        time.sleep(3)  
        move_s3_orginal_file(bucket_name, actual_file_name, existing_version, existing_file_name,app_id,env)
        move_s3_folder_reverse(bucket_name,actual_file_name, new_version, new_file_name,app_id,env)
        time.sleep(3)
        move_s3_orginal_file_reverse(bucket_name,actual_file_name, new_version, new_file_name,app_id,env)
        query_insert = f"""
        INSERT INTO {schema}.{file_log_table} (id, actual_file_name, active_file_name, done_by, action,app_id,env)
        VALUES (
            '{id_to_update}',
            '{actual_file_name}',
            '{new_file_name}',
            '{done_by}',
            'File Version Reverted',
            '{app_id}',
            '{env}'
        );
        """
        
        # Execute the INSERT query using the existing db_insert function
        db_insert(query_insert)
        
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        
        lambda_client = boto3.client('lambda',region_name = region_used)
        lambda_response = lambda_client.invoke(
            FunctionName = Injestion_trigger,                                                                                
            Payload = json.dumps({
                'file_name':new_file_name,
                'bucket_name':bucket_name,
                'trigger_action':'Put',
                'app_name':app_name
            }),   
            InvocationType = 'Event'
        )
        
        return {"status":"200","Message":"file version reverted successfully"}  
        
    except Exception as e:
        print("An exception occurred:",e)
        if id_to_update:
            update_query = f"""
            UPDATE {schema}.{file_metadata_Table}
            SET sync_status = 'FAILED'
            WHERE id = '{id_to_update}';
            """
            update_db(update_query)
        return {"status":"500","Message":"Failed to revert version"}

def revert_file_version(actual_file_name,new_file_name,new_version,done_by,app_id,env):
    try:
        id_to_update = None
        #Checking for inprogress:
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE active_file_name = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'INPROGRESS'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        result = db_select(query)
        
        if result:
            return {"status": "100", "Message": "FILE IS INPROGRESS"}
        
        #Checking for deleting
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'DELETING'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        result = db_select(query)
        if result:
            return {"status": "100", "Message": "FILE IS DELETING"}
            
        
        query_select = f"""
        SELECT * FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        
        # Execute the SELECT query using the existing db_select function
        response = db_select(query_select)
        existing_version = int(response[0][4])                                          
        existing_file_name = response[0][2]
        id_to_update = response[0][0]
        
        print(f"ACTUAL FILE : {actual_file_name}---EXISTING VERSION : {existing_version}---EXISTING FILE NAME : {existing_file_name}")
        move_s3_file(bucket_name, actual_file_name, existing_version, existing_file_name,app_id,env)
        current_time = datetime.now(timezone.utc).isoformat()
    
        #file_versions scan
        query_select = f"""
        SELECT count_in_kb
        FROM {schema}.{file_version_table}
        WHERE actual_file_name = '{actual_file_name}'
            AND delete_status = 0
            AND active_file_name = '{new_file_name}'
            AND version = {int(new_version)};
        """
        
        # Execute the SELECT query using the existing db_select function
        select_response = db_select(query_select)
        count_files = int(select_response[0][0])
        #update total_versions, active_version, active_file_name
        # Define the update query using an f-string
        query_update = f"""
        UPDATE {schema}.{file_metadata_Table}
        SET active_version = {new_version},
            active_file_name = '{new_file_name}',
            last_updated_time = '{current_time}',
            count_in_kb = {count_files},
            remain_count_kb = 0,
            sync_status = 'INPROGRESS',
            trigger_count = 0
        WHERE id = '{id_to_update}';
        """
        
        # Execute the UPDATE query using the existing update_db function
        update_response = update_db(query_update)
        print(f"ACTUAL FILE : {actual_file_name}---NEW VERSION : {new_version}---NEW FILE NAME : {new_file_name}")
        move_s3_file_reverse(bucket_name,actual_file_name,new_version,new_file_name,app_id,env)
        query_insert = f"""
        INSERT INTO {schema}.{file_log_table} (id, actual_file_name, active_file_name, done_by, action,app_id,env)
        VALUES (
            '{id_to_update}',
            '{actual_file_name}',
            '{new_file_name}',
            '{done_by}',
            'File Version Reverted',
            '{app_id}',
            '{env}'
        );
        """
        
        # Execute the INSERT query using the existing db_insert function
        db_insert(query_insert)
        
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        
        lambda_client = boto3.client('lambda',region_name = region_used)
        lambda_response = lambda_client.invoke(
            FunctionName = Injestion_trigger,                                                                                
            Payload = json.dumps({
                'file_name':new_file_name,
                'bucket_name':bucket_name,
                'trigger_action':'Put',
                'app_name':app_name
            }),   
            InvocationType = 'Event'
        )
        
        return {"status":"200","Message":"Version reverted succesfully. Ingestion started ..."}
        
    except Exception as e:
        print("An exception occurred:",e)
        if id_to_update:
            update_query = f"""
            UPDATE {schema}.{file_metadata_Table}
            SET sync_status = 'FAILED'
            WHERE id = '{id_to_update}';
            """
            update_db(update_query)
        return {"status":"500","Message":"Failed to revert version"}

# FILE REVERTING VERSION FUNCTION END

# START INGESTION FUNCTION
def start_ingestion_job(knowledgebaseid, datasourceid):
    aws_region=region_used                                                 
    client_bedrock_agent = boto3.client('bedrock-agent', region_name=region_used)
    retries = 0
    MAX_RETRIES = 150  # Maximum number of retries                    
    RETRY_DELAY = 2
    while retries < MAX_RETRIES:
        try:
            start_ingestion_response = client_bedrock_agent.start_ingestion_job(
                knowledgeBaseId=knowledgebaseid,
                dataSourceId=datasourceid,
            )
            ingestion_job_id = start_ingestion_response['ingestionJob']['ingestionJobId']
            return ingestion_job_id
        except Exception as e:
            # Catch all exceptions
            retries += 1
            print(f"An exception occurred: {e}")
            print(f"Retrying in {RETRY_DELAY} seconds... (Attempt {retries}/{MAX_RETRIES})")
            time.sleep(RETRY_DELAY)



# DELETE WHOLE FILE FUNCTION START

def delete_whole_csv_fn(actual_file_name,active_file_name,active_version,done_by,app_id,env):
    try:
        id_file_metadata = None
        #Checking for inprogress:
         #Checking for inprogress:
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE active_file_name = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'INPROGRESS'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        result = db_select(query)
        
        if result:
            return {"status": "100", "Message": "FILE IS INPROGRESS"}
        
        #Checking for deleting
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'DELETING'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        result = db_select(query)
        if result:
            return {"status": "100", "Message": "FILE IS DELETING"}
            
        #update file_metadata table
        query = f"""
        SELECT id
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND app_id = '{app_id}'
          AND env = '{env}'
        LIMIT 1;
        """
        
        result = db_select(query)
        print("FILE DATA : ",result)
        id_file_metadata = result[0][0]
        query = f"""
        UPDATE {schema}.{file_metadata_Table}
        SET sync_status = 'DELETING',
            trigger_count = 0
        WHERE id = '{id_file_metadata}';
        """
        
        response = update_db(query)
        print("UPDATE SYNC STATUS DELETINGG IN DB")
        #delete file in kb
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        object_key_prefix = f"{app_name}/Dev_Documents_kb/{active_file_name}/" 
        object_key = f"{app_name}/Dev_Documents_kbview/{active_file_name}"  
        s3_delete = s3_client.delete_object(Bucket=bucket_name, Key=object_key)
        delete_folder(bucket_name,object_key_prefix)
        #delete folder 
        folder_path = f"{app_name}/Dev_Documents/{actual_file_name}_versions/"
        delete_folder(bucket_name,folder_path)
    
        #update file_versions table
        query = f"""
            SELECT id
            FROM {schema}.{file_version_table}
            WHERE actual_file_name = '{actual_file_name}'
              AND delete_status = 0
              AND app_id = '{app_id}'
              AND env = '{env}';
            """
        result = db_select(query)
        for version_id in result:
            update_query = f"""
            DELETE from {schema}.{file_version_table}
            WHERE id = '{version_id[0]}';
            """
            update_db(update_query)
        print("All updates completed.")
        query_insert = f"""
        INSERT INTO {schema}.{file_log_table} (id, actual_file_name, active_file_name, done_by, action,app_id,env)
        VALUES (
            '{id_file_metadata}',
            '{actual_file_name}',
            '{active_file_name}',
            '{done_by}',
            'Whole File Deleted',
            '{app_id}',
            '{env}'
        );
        """
        # Execute the INSERT query using the existing db_insert function
        db_insert(query_insert)
        
        lambda_client = boto3.client('lambda',region_name = region_used)
        lambda_response = lambda_client.invoke(
            FunctionName = Injestion_trigger,                                                                                
            Payload = json.dumps({
                'file_name':active_file_name,
                'bucket_name':bucket_name,
                'trigger_action':'Delete',
                'app_name':app_name
            }),   
            InvocationType = 'Event'
        )
        
        return {"status":"200","Message":"File deleted successfully"}
        
    except Exception as e:
        print("An exception occurred : ",e)
        if id_file_metadata:
            update_query = f"""
                    UPDATE {file_metadata_Table}
                    SET sync_status = 'FAILED'
                    WHERE id = '{id_file_metadata}';
                    """
                    # Execute the update query
            update_db(update_query)
        return {"status":"500","Message":"Failed to delete the file"}   

def delete_whole_file_fn(actual_file_name,active_file_name,active_version,done_by,app_id,env):
    try:
        id_file_metadata = None
        #Checking for inprogress:
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE active_file_name = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'INPROGRESS'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        result = db_select(query)
        
        if result:
            return {"status": "100", "Message": "FILE IS INPROGRESS"}
        
        #Checking for deleting
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'DELETING'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        result = db_select(query)
        if result:
            return {"status": "100", "Message": "FILE IS DELETING"}
            
        #update file_metadata table
        query = f"""
        SELECT id
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0     
          AND app_id = '{app_id}'
          AND env = '{env}'
        LIMIT 1;
        """
        
        result = db_select(query)
        print("FILE DATA : ",result)
        id_file_metadata = result[0][0]
        query = f"""
        UPDATE {schema}.{file_metadata_Table}
        SET sync_status = 'DELETING',
            trigger_count = 0
        WHERE id = '{id_file_metadata}';
        """
        # Execute the UPDATE query
        update_db(query)
        #delete file in kb
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        object_key = f"{app_name}/Dev_Documents_kb/{active_file_name}"   
        object_key_metadata = f"{app_name}/Dev_Documents_kb/{active_file_name}.metadata.json"      
        s3_delete = s3_client.delete_object(Bucket=bucket_name, Key=object_key)
        s3_delete = s3_client.delete_object(Bucket=bucket_name,Key=object_key_metadata)         
        #delete folder 
        folder_path = f"{app_name}/Dev_Documents/{actual_file_name}_versions/"
        delete_folder(bucket_name,folder_path)
        #update file_versions table
        query = f"""
            SELECT id
            FROM {schema}.{file_version_table}
            WHERE actual_file_name = '{actual_file_name}'
              AND delete_status = 0;
            """
        result = db_select(query)
        for version_id in result:
            update_query = f"""
            DELETE from {schema}.{file_version_table}
            WHERE id = '{version_id[0]}';
            """
            update_db(update_query)
        print("All updates completed.")
        query_insert = f"""
        INSERT INTO {schema}.{file_log_table} (id, actual_file_name, active_file_name, done_by, action,app_id,env)
        VALUES (
            '{id_file_metadata}',
            '{actual_file_name}',
            '{active_file_name}',
            '{done_by}',
            'Whole File Deleted',
            '{app_id}',
            '{env}'   
        );
        """
        # Execute the INSERT query using the existing db_insert function
        db_insert(query_insert)
        
        lambda_client = boto3.client('lambda',region_name = region_used)
        lambda_response = lambda_client.invoke(
            FunctionName = Injestion_trigger,                                                                                      
            Payload = json.dumps({
                'file_name':active_file_name,
                'bucket_name':bucket_name,
                'trigger_action':'Delete',
                'app_name':app_name
            }),   
            InvocationType = 'Event'
        )
        
        return {"status":"200","Message":"File deleted successfully"}
    except Exception as e:
        print("An exception occurred:",e)
        if id_file_metadata:
            update_query = f"""
                    UPDATE {file_metadata_Table}
                    SET sync_status = 'FAILED'
                    WHERE id = '{id_file_metadata}';
                    """
                    # Execute the update query
            update_db(update_query)
        return {"status":"500","Message":"File deletion failed"}

# DELETE WHOLE FILE FUNCTION END


# DELETE FILE VERSION FUNCTION START

def delete_version_csv_fn(actual_file_name,active_file_name,active_version,done_by,app_id,env):    
    try:
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        
        # Execute the first query
        select_response = db_select(query)
        print("IN SELECT RESPONSE : ",select_response)
        id_file_metadata = select_response[0][0]
        #Checking for inprogress:
        # Construct the first SELECT query to check for INPROGRESS status
        query_inprogress = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'INPROGRESS'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        
        # Execute the first query
        response_inprogress = db_select(query_inprogress)
        print("IN PROGRESS RESPONSE : ",response_inprogress)
        
        # Check if any items are returned for INPROGRESS status
        if response_inprogress:
            return {"status": "100", "Message": "FILE IS INPROGRESS"}
        
        # Construct the second SELECT query to check for DELETING status
        query_deleting = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'DELETING'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        
        # Execute the second query
        response_deleting = db_select(query_deleting)
        
        # Check if any items are returned for DELETING status
        if response_deleting:
            return {"status": "100", "Message": "FILE IS DELETING"}
            
        #update the file_version table
        query_select = f"""
        SELECT id
        FROM {schema}.{file_version_table}
        WHERE actual_file_name = '{actual_file_name}'
          AND active_file_name = '{active_file_name}'
          AND version = {active_version}
          AND delete_status = 0
          AND app_id = '{app_id}'
          AND env = '{env}';
        """

        response = db_select(query_select)
        if response:
            id_value = response[0][0]
            query_update = f"""
            DELETE from {schema}.{file_version_table}
            WHERE id = '{id_value}';
            """
            update_db(query_update)
    
        #delete version folder in s3 
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        folder_path = f"{app_name}/Dev_Documents/{actual_file_name}_versions/{actual_file_name}_version{active_version}/"
        delete_folder(bucket_name,folder_path)      
        query_insert = f"""
        INSERT INTO {schema}.{file_log_table} (id, actual_file_name, active_file_name, done_by, action,app_id,env)
        VALUES (
            '{id_file_metadata}',
            '{actual_file_name}',
            '{active_file_name}',
            '{done_by}',
            'File version {active_version} Deleted',
            '{app_id}',
            '{env}'
        );
        """
        # Execute the INSERT query using the existing db_insert function
        db_insert(query_insert)
        return{"status":"200","Message":"Version deleted succesfully"}
    except Exception as e:
        print("An exception occurred",e)
        query_select = f"""
        SELECT id
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        response = db_select(query_select)
        if response:
            id_to_update = response[0][0]
            query_update = f"""
            UPDATE {schema}.{file_metadata_Table}
            SET sync_status = 'FAILED'
            WHERE id = '{id_to_update}';
            """
            update_db(query_update)
        return{"status":"500","Message":"Version deletion failed"}
        
def delete_version_file_fn(actual_file_name,active_file_name,active_version,done_by,app_id,env):
    try:
        query = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        
        # Execute the first query
        select_response = db_select(query)
        print("IN SELECT RESPONSE : ",select_response)
        id_file_metadata = select_response[0][0]
        #Checking for inprogress:
        query_inprogress = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'INPROGRESS'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        
        # Execute the first query
        response_inprogress = db_select(query_inprogress)
        # Check if any items are returned for INPROGRESS status
        if response_inprogress:
            return {"status": "100", "Message": "FILE IS INPROGRESS"}
        
        # Checking for deleting
        query_deleting = f"""
        SELECT *
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND sync_status = 'DELETING'
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        
        # Execute the second query
        response_deleting = db_select(query_deleting)
        
        # Check if any items are returned for DELETING status
        if response_deleting:
            return {"status": "100", "Message": "FILE IS DELETING"}
               
        #update the file_version table
        query_select = f"""
        SELECT id
        FROM {schema}.{file_version_table}
        WHERE actual_file_name = '{actual_file_name}'
          AND active_file_name = '{active_file_name}'
          AND version = {active_version}
          AND delete_status = 0
          AND app_id = '{app_id}'
          AND env = '{env}';
        """

        response = db_select(query_select)
        if response:
            id_value = response[0][0]
            query_update = f"""
            DELETE from {schema}.{file_version_table}
            WHERE id = '{id_value}';
            """
            update_db(query_update)
    
        #delete version folder in s3 
        app_query = f"SELECT app_name from {schema}.{app_metadata_table} WHERE app_id = '{app_id}';"
        app_name_result = db_select(app_query)
        app_name = app_name_result[0][0]
        folder_path = f"{app_name}/Dev_Documents/{actual_file_name}_versions/{actual_file_name}_version{active_version}/"
        delete_folder(bucket_name,folder_path)      
        query_insert = f"""
        INSERT INTO {schema}.{file_log_table} (id, actual_file_name, active_file_name, done_by, action)
        VALUES (
            '{id_file_metadata}',
            '{actual_file_name}',
            '{active_file_name}',
            '{done_by}',
            'File version {active_version} Deleted'
        );
        """
        # Execute the INSERT query using the existing db_insert function
        db_insert(query_insert)
        return{"status":"200","Message":"Version deleted succesfully"}
    
    except Exception as e:
        print("An exception occurred :",e)
        query_select = f"""
        SELECT id
        FROM {schema}.{file_metadata_Table}
        WHERE actual_filename = '{actual_file_name}'
          AND delete_status = 0
          AND app_id = '{app_id}'
          AND env = '{env}';
        """
        response = db_select(query_select)
        if response:
            id_to_update = response[0][0]
            query_update = f"""
            UPDATE {schema}.{file_metadata_Table}
            SET sync_status = 'FAILED'
            WHERE id = '{id_to_update}';
            """
            update_db(query_update)
        return{"status":"500","Message":"Version deletion failed"} 
        
def delete_folder(bucket_name, folder_path):
    try:
        # List objects in the specified folder
        #response = s3_client.list_objects_v2(Bucket=bucket_name, Prefix=folder_path)
        paginator = s3_client.get_paginator('list_objects_v2')
        page_iterator = paginator.paginate(Bucket=bucket_name,  Prefix=folder_path)  
        for page in page_iterator:
            # Delete each object in the folder
            if 'Contents' in page:
                for obj in page['Contents']:
                    key = obj['Key']
                    s3_client.delete_object(Bucket=bucket_name, Key=key)
                    print(f"Deleted object: {key}")
    
        # Delete the folder-like prefix itself
        s3_client.delete_object(Bucket=bucket_name, Key=folder_path)
    
        print(f"Deleted folder: {folder_path}")
    except Exception as e:
        print("An exception occurred at function delete_folder: ",e)   

# DELETE FILE VERSION FUNCTION END
def lambda_handler(event, context):
    # TODO implement
    print("EVENT : ",event)
    event_type = event["event_type"]
    app_id = event["app_id"]
    env = "Development"   
    
    # NEW FILE UPLOAD API
    if event_type == 'new_file_upload':  
        # file_content = event['file_content']
        # print(file_content)
        input_file_key = event['input_file_key']                                                                     
        file_name = event['file_name']
        file_type = file_name.split('.')[-1] 
        print('file_type',file_type)                        
        access_type = event['access_type']
        uploaded_by = event['uploaded_by']
        # content_decoded = base64.b64decode(file_content)
        # print(content_decoded,'content_decoded')
    
        if file_type != 'csv' and file_type != 'xlsx' and file_type != 'docx' and file_type != 'pdf':
            print("failed") 
            return {"status":"500","Message":"Invalid file type"}
    
        if file_type == 'csv' or file_type =='xlsx':  
            print('new csv upload invoked')
            result = new_csv_upload_process(file_name,file_type,access_type,uploaded_by,input_file_key,app_id,env)
        
        else:
            print('new file upload invoked')
            result = new_file_upload_process(file_name,file_type,access_type,uploaded_by,input_file_key,app_id,env)
        
        return result    
        
        
    # NEW VERSION UPLOAD API
    if event_type == 'new_version_file_upload': 
        # file_content = event['file_content']
        input_file_key = event['input_file_key']
        actual_file_name = event['actual_file_name']
        file_name = event['file_name']
        file_type = file_name.split('.')[-1]                                                             
        access_type = event['access_type']
        uploaded_by = event['uploaded_by']
        # content_decoded = base64.b64decode(file_content)
    
        if file_type != 'csv' and file_type != 'xlsx' and file_type != 'docx' and file_type != 'pdf':
            print("failed") 
            return {"status":"500","Message":"Invalid file type"}
    
        if file_type == 'csv' or file_type == 'xlsx' :  
            print("NEW CSV VERSION UPLOAD INVOKE")
            result = new_version_csv_upload_process(actual_file_name,file_name,file_type,access_type,uploaded_by,input_file_key,app_id,env)
        
        else:
            result = new_version_file_upload_process(actual_file_name,file_name,file_type,access_type,uploaded_by,input_file_key,app_id,env)
        
        return result  
        
    # FILE PRE PROCESSING API
    if event_type == 'html_word_preprocess_content':  
        file_content = event['file_content']
        print("EVENT : ",event)             
        print(file_content)                                                                       
        file_name = event['file_name']
        print("FILE NAME",file_name)
        uploaded_by = event['uploaded_by']
        file_type = file_name.split('.')[-1]  
        print("FILE TYPE : ",file_type)
        
        query = f"""
                SELECT * FROM {schema}.{preprocess_table}
                WHERE input_filename = '{file_name}' AND delete_status = 0 AND app_id = '{app_id}' AND env = '{env}';
                """  
        print("CHECKING IF FILE EXISTS")
        response = db_select(query)
        print("FILE NAME CHECK RDS RESPONSE",response)
        if len(response) != 0:
            return {"status":"500","Message":"FILE NAME ALREADY EXISTS"} 
        if file_type != 'csv' and file_type != 'xlsx' and file_type != 'docx':
            print("failed") 
            return {"status":"500","Message":"Invalid file type"}
        if file_type == 'docx':
            doc =Document()
            print("IN WORD PRE PROCESSING")         
            result = word_pre_processing(file_content,file_type,uploaded_by,file_name,app_id,env)
        return result  
        
    # REVERTING FILE VERSION API
    if event_type == 'revert_version':
        actual_file_name = event['actual_file_name']
        new_file_name =  event['new_file_name']
        new_version = event['new_version']
        done_by = event["uploaded_by"]
        file_type =  new_file_name.split('.')[-1]  
    
        if file_type != 'csv' and file_type != 'xlsx' and file_type != 'docx' and file_type != 'pdf':
            print("failed") 
            return {"status":"500","Message":"Invalid file type"}
    
        if file_type =='csv' or file_type == 'xlsx' :
            result = revert_csv_version(actual_file_name,new_file_name,new_version,done_by,app_id,env)     
        
        else:
            result = revert_file_version(actual_file_name,new_file_name,new_version,done_by,app_id,env)
        
        return result 
    
    # DELETE WHOLE FUNCTION API
    if event_type == 'delete_whole_file':
        actual_file_name = event['actual_file_name']
        active_file_name = event['active_file_name']    
        active_version = event['active_version']
        done_by = event["uploaded_by"]
        file_type =  active_file_name.split('.')[-1]
    
        if file_type != 'csv' and file_type != 'xlsx' and file_type != 'docx' and file_type != 'pdf':
            print("failed") 
            return {"status":"500","Message":"Invalid file type"}
    
        if file_type =='csv' or file_type == 'xlsx':
            result = delete_whole_csv_fn(actual_file_name,active_file_name,active_version,done_by,app_id,env)
        
        else:
            result = delete_whole_file_fn(actual_file_name,active_file_name,active_version,done_by,app_id,env)
        
        return result
        
    # DELETE VERSION API
    if event_type == 'delete_version':  
        actual_file_name = event['actual_file_name']
        active_file_name = event['active_file_name']
        delete_version = int(event['delete_version'])
        done_by = event["uploaded_by"]
        file_type =  active_file_name.split('.')[-1]     
    
        if file_type != 'csv' and file_type != 'xlsx' and file_type != 'docx' and file_type != 'pdf':
            print("failed") 
            return {"status":"500","Message":"Invalid file type"}   
    
        if file_type =='csv' or file_type == 'xlsx' :  
            result = delete_version_csv_fn(actual_file_name,active_file_name,delete_version,done_by,app_id,env)
        
        else:
            result = delete_version_file_fn(actual_file_name,active_file_name,delete_version,done_by,app_id,env)
        
        return result     

    #CHECK FILE AVAILABILITY API
    if event_type == 'check_file_availability':
        app_id = event['app_id']
        file_name = event['file_name']
        env = event['env']
        query = f"""
            SELECT * FROM {schema}.{file_version_table}
            WHERE active_file_name = '{file_name}' AND delete_status = 0 AND app_id = '{app_id}' AND env = '{env}';
            """  
        print("CHECKING IF FILE EXISTS")
        response = db_select(query)
        print("FILE NAME CHECK RDS RESPONSE",response)
        if len(response) != 0:
            print("FILE NAME ALREADY EXISTS")
            return {"status":"500","Message":"FILE NAME ALREADY EXISTS"}
        else:
            print("FILE DOES NOT EXIST THIS FILE CAN BE INSERTED")
            #Checking for inprogress:
            query = f"""
            SELECT *
            FROM {schema}.{file_metadata_Table}
            WHERE active_file_name = '{file_name}'
            AND delete_status = 0
            AND sync_status = 'INPROGRESS'
            AND app_id = '{app_id}' AND env = '{env}';
            """
            
            result = db_select(query)
            print("IN PROGRESS QUERY RESULT",result)
            
            if result:
                return {"status": "100", "Message": "FILE IS INPROGRESS"}
            
            #Checking for deleting
            query = f"""
            SELECT *
            FROM {schema}.{file_metadata_Table}
            WHERE actual_filename = '{file_name}'
            AND delete_status = 0
            AND sync_status = 'DELETING'
            AND app_id = '{app_id}' AND env = '{env}';
            """
            
            result = db_select(query)
            print("DELETING QUERY RESULT : ",result)

            if result:
                return {"status": "100", "Message": "FILE IS DELETING"}
         
            return {"status":"200","Message":"FILE DOES NOT EXIST THIS FILE CAN BE INSERTED"}     